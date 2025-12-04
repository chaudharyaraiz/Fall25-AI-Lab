from flask import Flask, render_template, request, redirect, url_for, send_file
from werkzeug.utils import secure_filename
import pandas as pd
from io import BytesIO
from docx import Document
import os

UPLOAD_FOLDER = "uploads"
ALLOWED_EXTENSIONS = {"csv"}

app = Flask(__name__)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        file = request.files.get("file")
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
            file.save(path)
            df = pd.read_csv(path)
            preview = df.head(100).to_html(classes="table table-striped table-bordered", index=False, escape=False)
            request.environ["uploaded_path"] = path
            return render_template("result.html", table_html=preview, filename=filename)
        return redirect(url_for("index"))
    return render_template("index.html")

@app.route("/download/<filename>")
def download(filename):
    path = os.path.join(app.config["UPLOAD_FOLDER"], secure_filename(filename))
    if not os.path.exists(path):
        return redirect(url_for("index"))
    df = pd.read_csv(path)
    doc = Document()
    doc.add_heading(filename, level=2)
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, item in enumerate(row):
            row_cells[i].text = "" if pd.isna(item) else str(item)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return send_file(bio, as_attachment=True, download_name=f"{os.path.splitext(filename)[0]}.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == "__main__":
    app.run(debug=True)
